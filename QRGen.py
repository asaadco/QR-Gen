import docx
import os
import pandas as pd
import math
import string
from datetime import datetime
import time
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from face_crop import face_crop
from fuzzywuzzy import process
def fexists(innovator, media, Fuzzy):
    if(media):

        if(Fuzzy):
            return process.extractOne(innovator, media)[1] > 50
        else:
            return innovator in media
    else: return False

def fmatch(innovator, media, Fuzzy):
    if(media):
        if(Fuzzy):
            print(process.extractOne(innovator, media))
            return process.extractOne(innovator, media)[0]
        else:
            return media.index(innovator)
    else: return



# This function is used for adding the images into a table
def tadd_Photo(text, pic, table, row, col, face):
    if(face):
        pic = face_crop(pic)
    pt = table.cell(row, col).paragraphs[0]
    pt.text = ''
    #pt.style = document.styles['No Spacing']
    pt.space_before = 0
    pt.space_after = 0
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pt.add_run()
    try:
        r.add_picture(pic,width=Inches(1.0), height=Inches(1.0))
    except:
        r.add_picture("stock.jpg", width=Inches(1.0), height=Inches(1.0))
    r.add_break()
    r.add_text(text)

def generateDoc(dbdir, attachfolder, Fuzzy, Face):


    df = pd.read_excel(dbdir, "Main") # Reading from the excel file 
    ideasDoc = Document() # Create a new Word Document
    storiesDoc = Document()
    #styles = document.styles # Save the styles available by default in any Word Document into styles variable for future use
    # print("Column headings:")


    print(df.columns)

    # To filter the rows based on Submission Type

    print(df)
    print(len(df))
    for index, row in df.iterrows():
        ###############PREPROCESSING###################
        uncleanDepts = set(row["Departments"].split(";"))
        departments = set()
        for dept in uncleanDepts:
            departments.add(string.capwords(dept.split("/")[0].strip()))
        departments = "\n".join(departments)
        departments = departments.replace("\n", "", 1)
        media = list()
        mediaPath = attachfolder+"\\"+str(row["ID"])+"\\"
        if(os.path.exists(mediaPath)):
            media = os.listdir(mediaPath)
        ######## Initialize Variables ################
        innovatorsRaw = row["Inventor"].split(";")
        innovatorsRaw.remove("")
        innovators = list()
        for innovator in innovatorsRaw:
            tmp = innovator.split(",")
            name = tmp[1].split(" ")[1] + " " + tmp[0]
            innovators.append(string.capwords(name))
        #                  "AlGhamdi, Asaad S;".Split(;) => {"AlGhamdi, Asaad S, ", ""}
        cleanA = list()
        for attachment in media:
            cleanA.append(string.capwords(attachment.split(".")[0]))
            
    
        rows = math.ceil(len(innovators)/2)
        col = 3
        print("---------------------------------")
        print("Processing ... ID: #", row["ID"])
        print(row["Title"])
        print("# of Innovators:", len(innovators))
        print(innovators)
        print("Departments", departments)
        print("Attachments", media)
        print("Clean Attachments", cleanA)
        print("---------------------------------")

    

    ###################################


    #    p = document.add_paragraph(row["Body"])

    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True

    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')

    # document.add_paragraph(
    #     'first item in unordered list', style='List Bullet'
    # )
    # document.add_paragraph(
    #     'first item in ordered list', style='List Number'
    # )
    # document.add_picture(r'Attachments\12\Lahouari_Ghouti[1].jpg', width=Inches(1.25))


        # Need to make the rows and columns into variables
        if(row['Submission Type'] == "Idea"):
            table = ideasDoc.add_table(rows=rows, cols=3)
            table.style = 'TableGrid'
            table.autofit = False
            k = 0
            for i in range(0, rows):
                table.cell(i,0).width = Inches(1.25)
                table.cell(i,1).width = Inches(1.25)
                print("CURRENTLY PROCESSING:", innovators[k])
                
                if(k < len(innovators)):
                    #if(innovators[k] in cleanA):
                    if(fexists(innovators[k], cleanA, Fuzzy)):
                        # We want to use Fuzzy Matching instead of cleanA.index(innovator[k]). Define a method
                        tadd_Photo(innovators[k], mediaPath+"\\"+media[cleanA.index(fmatch(innovators[k], cleanA, Fuzzy))], table, i, 0, Face)
                    else:
                        tadd_Photo(innovators[k], "stock.jpg", table, i, 0, False)               
                    k = k + 1
                if(k < len(innovators)):
                    if(fexists(innovators[k], cleanA, Fuzzy)):
                    #if(innovators[k] in cleanA):
                        tadd_Photo(innovators[k], mediaPath+"\\"+media[cleanA.index(fmatch(innovators[k], cleanA, Fuzzy))], table, i, 1, Face)
                    else: 
                        tadd_Photo(innovators[k], "stock.jpg", table, i, 1, False)
                    k = k + 1
                else:
                    table.cell(i, 0).merge(table.cell(i, 1))
                if((i+1) != rows):
                    table.cell(i, 2).merge(table.cell(i+1, 2))
            pt = table.cell(0, 2).paragraphs[0]
            table.cell(0, 2).width = Inches(3.70)
            pt.text = row["Title"]+"\n"+departments+"\n\n" + row["Body"] + "\n\n" 
            # row["Title"] + "\n" + "\n\n" + row["Body"] + "\n"
            if("Photo" in cleanA):
                r = pt.add_run()
                r.add_picture(mediaPath+"\\"+media[cleanA.index(fmatch("Photo", cleanA, Fuzzy))], width=Inches(3.70), height=Inches(2.63))    
            ideasDoc.add_page_break()
        else:
            table = storiesDoc.add_table(rows=1, cols=1)
            table.style = 'TableGrid'
            table.autofit = False
            pt = table.cell(0, 0).paragraphs[0]
            pt.text = row["Title"]+"\n"+departments+"\n\n"+ row["Body"] + "\n\n"
            if("Photo" in cleanA):
                r = pt.add_run()
                r.add_picture(mediaPath+"\\"+media[cleanA.index(fmatch("Photo", cleanA, Fuzzy))], width=Inches(5.77), height=Inches(2.63))
            storiesDoc.add_page_break()




        # tadd_Photo("Lahouari Ghouti", 'Lahouari_Ghouti[1].jpg', table, 0, 1)

        # tadd_Photo("Lahouari Ghouti", 'Lahouari_Ghouti[1].jpg', table, 1, 0)

        # tadd_Photo("Lahouari Ghouti", 'Lahouari_Ghouti[1].jpg', table, 1, 1)


        # This is how you adjust the width of cells
        # table.cell(0,0).width = Inches(.50)
        # table.cell(0,1).width = Inches(.50)
        # table.cell(1,1).width = Inches(.50)
        # table.cell(1,0).width = Inches(.50)
        # table.cell(1,0).width = Inches(.50)
        # table.cell(1,1).width = Inches(.50)


        ## table_cell(i, 2).merge(table.cell(i+1, 2))
        # table.cell(0, 2).merge(table.cell(1, 2))

        
    # if(index+1 != len(df)):
    #     document.add_page_break()
    ideasDoc.save("QR Ideas"+ ' '+datetime.strftime(datetime.now(), '%b %d %Y')+' .docx')
    storiesDoc.save("QR Stories"+ ' '+datetime.strftime(datetime.now(), '%b %d %Y')+' .docx')


